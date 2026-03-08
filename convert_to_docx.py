#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
将食材包订阅平台使用说明书Markdown文件转换为Word文档(docx)
要求：安装python-docx库 (pip install python-docx)
"""

import re
import sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

def create_document():
    """创建Word文档并设置样式"""
    doc = Document()

    # 设置默认字体
    styles = doc.styles
    normal_style = styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Times New Roman'
    normal_font.size = Pt(12)

    # 设置中文字体
    normal_font._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 设置段落格式
    normal_paragraph_format = normal_style.paragraph_format
    normal_paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal_paragraph_format.line_spacing = 1.2  # 1.2倍行距
    normal_paragraph_format.space_after = Pt(6)  # 段后0.5行 (12pt * 0.5 = 6pt)
    normal_paragraph_format.first_line_indent = Pt(24)  # 首行缩进2字符 (12pt * 2 = 24pt)

    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)      # 2.54厘米
        section.bottom_margin = Inches(1.0)   # 2.54厘米
        section.left_margin = Inches(1.25)    # 3.18厘米
        section.right_margin = Inches(1.25)   # 3.18厘米

    # 创建标题样式
    # 标题1
    if 'Heading 1' not in styles:
        heading1_style = styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)
    else:
        heading1_style = styles['Heading 1']
    heading1_font = heading1_style.font
    heading1_font.name = 'Times New Roman'
    heading1_font.size = Pt(24)
    heading1_font.bold = True
    heading1_font._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    heading1_paragraph_format = heading1_style.paragraph_format
    heading1_paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading1_paragraph_format.space_before = Pt(24)
    heading1_paragraph_format.space_after = Pt(18)

    # 标题2
    if 'Heading 2' not in styles:
        heading2_style = styles.add_style('Heading 2', WD_STYLE_TYPE.PARAGRAPH)
    else:
        heading2_style = styles['Heading 2']
    heading2_font = heading2_style.font
    heading2_font.name = 'Times New Roman'
    heading2_font.size = Pt(18)
    heading2_font.bold = True
    heading2_font._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    heading2_paragraph_format = heading2_style.paragraph_format
    heading2_paragraph_format.space_before = Pt(18)
    heading2_paragraph_format.space_after = Pt(12)

    # 标题3
    if 'Heading 3' not in styles:
        heading3_style = styles.add_style('Heading 3', WD_STYLE_TYPE.PARAGRAPH)
    else:
        heading3_style = styles['Heading 3']
    heading3_font = heading3_style.font
    heading3_font.name = 'Times New Roman'
    heading3_font.size = Pt(14)
    heading3_font.bold = True
    heading3_font._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    heading3_paragraph_format = heading3_style.paragraph_format
    heading3_paragraph_format.space_before = Pt(12)
    heading3_paragraph_format.space_after = Pt(6)

    # 标题4
    if 'Heading 4' not in styles:
        heading4_style = styles.add_style('Heading 4', WD_STYLE_TYPE.PARAGRAPH)
    else:
        heading4_style = styles['Heading 4']
    heading4_font = heading4_style.font
    heading4_font.name = 'Times New Roman'
    heading4_font.size = Pt(12)
    heading4_font.bold = True
    heading4_font._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    heading4_paragraph_format = heading4_style.paragraph_format
    heading4_paragraph_format.space_before = Pt(9)
    heading4_paragraph_format.space_after = Pt(3)

    return doc

def parse_markdown(md_content):
    """解析Markdown内容，返回结构化的段落列表"""
    lines = md_content.split('\n')
    paragraphs = []

    i = 0
    while i < len(lines):
        line = lines[i].strip()

        # 空行
        if not line:
            paragraphs.append({'type': 'empty'})
            i += 1
            continue

        # 标题
        if line.startswith('# '):
            paragraphs.append({'type': 'heading1', 'text': line[2:].strip()})
            i += 1
            continue
        elif line.startswith('## '):
            paragraphs.append({'type': 'heading2', 'text': line[3:].strip()})
            i += 1
            continue
        elif line.startswith('### '):
            paragraphs.append({'type': 'heading3', 'text': line[4:].strip()})
            i += 1
            continue
        elif line.startswith('#### '):
            paragraphs.append({'type': 'heading4', 'text': line[5:].strip()})
            i += 1
            continue

        # 列表项
        if line.startswith('- ') or line.startswith('* '):
            items = []
            while i < len(lines) and (lines[i].startswith('- ') or lines[i].startswith('* ') or lines[i].startswith('  ')):
                items.append(lines[i].strip())
                i += 1
            paragraphs.append({'type': 'bullet_list', 'items': items})
            continue

        # 有序列表
        match = re.match(r'^(\d+)\.\s+(.+)$', line)
        if match:
            items = []
            while i < len(lines) and re.match(r'^(\d+)\.\s+(.+)$', lines[i].strip()):
                items.append(lines[i].strip())
                i += 1
            paragraphs.append({'type': 'ordered_list', 'items': items})
            continue

        # 表格（简化处理）
        if '|' in line and '---' not in line:
            # 收集表格行
            table_rows = []
            while i < len(lines) and '|' in lines[i]:
                table_rows.append(lines[i].strip())
                i += 1
            # 跳过分隔行
            if i < len(lines) and '---' in lines[i]:
                i += 1
                while i < len(lines) and '|' in lines[i]:
                    table_rows.append(lines[i].strip())
                    i += 1
            paragraphs.append({'type': 'table', 'rows': table_rows})
            continue

        # 代码块
        if line.startswith('```'):
            # 跳过代码块
            i += 1
            while i < len(lines) and not lines[i].startswith('```'):
                i += 1
            if i < len(lines) and lines[i].startswith('```'):
                i += 1
            continue

        # 普通段落
        text_lines = []
        while i < len(lines) and lines[i].strip() and not (
            lines[i].startswith('#') or
            lines[i].startswith('- ') or
            lines[i].startswith('* ') or
            re.match(r'^\d+\.\s', lines[i].strip()) or
            '|' in lines[i] or
            lines[i].startswith('```')
        ):
            text_lines.append(lines[i].strip())
            i += 1

        if text_lines:
            paragraphs.append({'type': 'paragraph', 'text': ' '.join(text_lines)})
            continue

        # 如果都不匹配，跳过这一行
        i += 1

    return paragraphs

def add_content_to_document(doc, paragraphs):
    """将解析后的段落添加到Word文档"""
    for para_info in paragraphs:
        if para_info['type'] == 'empty':
            doc.add_paragraph()  # 空行
        elif para_info['type'] == 'heading1':
            doc.add_heading(para_info['text'], level=1)
        elif para_info['type'] == 'heading2':
            doc.add_heading(para_info['text'], level=2)
        elif para_info['type'] == 'heading3':
            doc.add_heading(para_info['text'], level=3)
        elif para_info['type'] == 'heading4':
            doc.add_heading(para_info['text'], level=4)
        elif para_info['type'] == 'paragraph':
            p = doc.add_paragraph(para_info['text'])
            # 确保段落格式
            p.style = doc.styles['Normal']
        elif para_info['type'] == 'bullet_list':
            for item in para_info['items']:
                # 移除列表标记
                item_text = re.sub(r'^[-*]\s+', '', item)
                p = doc.add_paragraph(item_text, style='List Bullet')
        elif para_info['type'] == 'ordered_list':
            for item in para_info['items']:
                # 移除数字标记
                item_text = re.sub(r'^\d+\.\s+', '', item)
                p = doc.add_paragraph(item_text, style='List Number')
        elif para_info['type'] == 'table':
            rows = para_info['rows']
            if len(rows) > 0:
                # 解析表格
                table_data = []
                for row in rows:
                    cells = [cell.strip() for cell in row.split('|') if cell.strip()]
                    if cells:
                        table_data.append(cells)

                if table_data:
                    # 创建表格
                    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                    table.style = 'Table Grid'

                    for i, row_cells in enumerate(table_data):
                        row = table.rows[i]
                        for j, cell_text in enumerate(row_cells):
                            if j < len(row.cells):
                                row.cells[j].text = cell_text

def main():
    """主函数"""
    input_file = '食材包订阅平台使用说明书.md'
    output_file = '食材包订阅平台使用说明书.docx'

    try:
        # 读取Markdown文件
        print(f"正在读取文件: {input_file}")
        with open(input_file, 'r', encoding='utf-8') as f:
            md_content = f.read()

        # 创建文档
        print("正在创建Word文档...")
        doc = create_document()

        # 解析Markdown
        print("正在解析Markdown内容...")
        paragraphs = parse_markdown(md_content)

        # 添加内容
        print("正在添加内容到Word文档...")
        add_content_to_document(doc, paragraphs)

        # 保存文档
        print(f"正在保存到: {output_file}")
        doc.save(output_file)

        print(f"转换完成！文档已保存为: {output_file}")
        print("\n文档设置：")
        print("- 中文字体：宋体")
        print("- 英文字体：Times New Roman")
        print("- 页边距：上下2.54厘米，左右3.18厘米")
        print("- 行距：1.2倍")
        print("- 段落：段后0.5行，首行缩进2字符")

    except FileNotFoundError:
        print(f"错误：找不到文件 {input_file}")
        print("请确保Markdown文件与脚本在同一目录下")
    except Exception as e:
        print(f"转换过程中发生错误: {e}")
        import traceback
        traceback.print_exc()

if __name__ == '__main__':
    main()